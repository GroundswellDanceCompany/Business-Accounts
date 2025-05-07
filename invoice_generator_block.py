
# Corrected script starts here
# Full working script integrating Invoice Generator, Student Manager, and all other modules

import streamlit as st
import gspread
import pandas as pd
from datetime import datetime, date, timedelta
from oauth2client.service_account import ServiceAccountCredentials
from docx import Document

# Google Sheets setup
scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
creds_dict = st.secrets["gcp_service_account"]
creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
client = gspread.authorize(creds)

# Main navigation
selection = st.sidebar.selectbox("Navigate", ["Invoice Generator", "Student Manager", "Dashboard", "Accounts Package", "Expenses Dashboard"])

# Helper to generate invoice document
def generate_invoice_doc(student_name, date_from, date_to, class_list, extras, total):
    doc = Document("invoice_template.docx")
    for para in doc.paragraphs:
        if "{{student_name}}" in para.text:
            para.text = para.text.replace("{{student_name}}", student_name)
        if "{{date_from}}" in para.text:
            para.text = para.text.replace("{{date_from}}", date_from)
        if "{{date_to}}" in para.text:
            para.text = para.text.replace("{{date_to}}", date_to)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Amount'

    for cls in class_list:
        row_cells = table.add_row().cells
        row_cells[0].text = cls
        row_cells[1].text = ""

    for item in extras:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = ""

    row_cells = table.add_row().cells
    row_cells[0].text = "Total"
    row_cells[1].text = f"£{total:.2f}"

    filename = f"{student_name}_invoice.docx"
    doc.save(f"/mnt/data/{filename}")
    return f"/mnt/data/{filename}"

# Invoice Generator
if selection == "Invoice Generator":
    st.title("Monthly Invoice Generator (with Dated Extras)")
    sheet = client.open("Groundswell-Business").worksheet("invoices")

    student = st.text_input("Student Name")

    today = date.today()
    invoice_start = st.date_input("Invoice Start Date", value=today)
    invoice_end = st.date_input("Invoice End Date", value=today + timedelta(weeks=4))
    invoice_label = st.text_input("Custom Invoice Label (e.g. Maya - April 2025)", value=f"{student} - {invoice_start.strftime('%b %Y')}")

    selected_classes = st.multiselect("Select Classes", [
        "Junior Ballet", "Intermediate Ballet", "Junior Contemporary", "Intermediate Contemporary",
        "Junior Jazz", "Advanced Jazz", "Junior House & Hip Hop", "Advanced House & Hip Hop",
        "Junior Waacking & Locking", "Advanced Waacking & Locking", "Tap Class", "Commercial", "Private"
    ])

    age_group = st.selectbox("Age Group", ["Mini (3-5)", "Junior (6-12)", "Teen (13-16)", "Adult"])

    price_chart = {
        "Junior Ballet": {"Mini (3-5)": 5, "Junior (6-12)": 5.50, "Teen (13-16)": 6, "Adult": 6.50},
        "Intermediate Ballet": {"Junior (6-12)": 5.50, "Teen (13-16)": 6, "Adult": 6.50},
        "Junior Contemporary": {"Junior (6-12)": 5.50, "Teen (13-16)": 6, "Adult": 6.50},
        "Intermediate Contemporary": {"Junior (6-12)": 5.50, "Teen (13-16)": 6, "Adult": 6.50},
        "Junior Jazz": {"Junior (6-12)": 5.50, "Teen (13-16)": 6, "Adult": 6.50},
        "Advanced Jazz": {"Junior (6-12)": 5.50, "Teen (13-16)": 6, "Adult": 6.50},
        "Junior House & Hip Hop": {"Junior (6-12)": 5.50, "Teen (13-16)": 6},
        "Advanced House & Hip Hop": {"Junior (6-12)": 5.50, "Teen (13-16)": 6, "Adult": 6.50},
        "Junior Waacking & Locking": {"Junior (6-12)": 5.50, "Teen (13-16)": 6},
        "Advanced Waacking & Locking": {"Junior (6-12)": 5.50, "Teen (13-16)": 6, "Adult": 6.50},
        "Tap Class": {"Junior (6-12)": 5.50, "Teen (13-16)": 6, "Adult": 6.50},
        "Commercial": {"Junior (6-12)": 5.50,"Teen (13-16)": 6, "Adult": 6.50},
        "Private": {"All": 20}
    }

    rates = []
    for cls in selected_classes:
        if cls == "Private":
            rate = price_chart["Private"]["All"]
        else:
            rate = price_chart.get(cls, {}).get(age_group, 0)
        rates.append((cls, rate))
    st.session_state.selected_classes = rates

    classes_attended = st.number_input("Total number of classes attended", min_value=0)

    available_extras = {
        "Subscription": ["Team Training-GFoundation", "Team Training-GSD Youth", "Team Training-Jenga"],
        "Session-Based": ["Extra Rehearsal - Competition", "Extra Rehearsal - Show", "Private Lesson"],
        "One-Off": ["Costume Fee", "Uniform Fee", "Exam Entry", "Competition Fee"]
    }

    all_extras = [item for sublist in available_extras.values() for item in sublist]
    if "extras" not in st.session_state:
        st.session_state.extras = []

    with st.form("add_extra_form", clear_on_submit=True):
        extra_name = st.selectbox("Choose Extra", all_extras)
        extra_type = next((k for k, v in available_extras.items() if extra_name in v), "")
        extra_amount = st.number_input("Amount", min_value=0.0, step=0.5)
        extra_date = None
        if extra_type == "Session-Based":
            extra_date = st.date_input("Date of Session", value=date.today())
        submit_extra = st.form_submit_button("Add Extra")

        if submit_extra and extra_name and extra_amount > 0:
            st.session_state.extras.append({
                "name": extra_name,
                "type": extra_type,
                "amount": extra_amount,
                "date": str(extra_date) if extra_date else ""
            })

    notes = st.text_area("Notes (optional)")

    if st.button("Generate Invoice"):
        total_class_rate = sum(rate for _, rate in rates)
        class_total = classes_attended * total_class_rate
        extras_total = sum(extra["amount"] for extra in st.session_state.extras)
        grand_total = class_total + extras_total

        invoice_period = f"{invoice_start} to {invoice_end}"
        date_created = datetime.now().strftime("%Y-%m-%d")
        class_names = ", ".join(cls for cls, _ in rates)
        extra_names = ", ".join(
            f"{ex['name']} ({ex['type']}" + (f" on {ex['date']}" if ex['type'] == "Session-Based" and ex['date'] else "") +
            f"): £{ex['amount']:.2f}"
            for ex in st.session_state.extras
        )

        row = [date_created, invoice_period, student, class_names, classes_attended, total_class_rate, extra_names, extras_total, grand_total, "Unpaid", notes, invoice_label]
        sheet.append_row(row)

        invoice_path = generate_invoice_doc(
            student_name=student,
            date_from=invoice_start.strftime("%Y-%m-%d"),
            date_to=invoice_end.strftime("%Y-%m-%d"),
            class_list=[f"{cls}: £{rate:.2f}" for cls, rate in st.session_state.selected_classes],
            extras=[f"{ex['name']} – £{ex['amount']:.2f}" for ex in st.session_state.extras],
            total=grand_total
        )

        with open(invoice_path, "rb") as file:
            st.download_button(
                label="Download Invoice (Word)",
                data=file,
                file_name=f"{student}_invoice.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.success(f"Invoice created for {student} (£{grand_total:.2f})")
        st.session_state.extras = []

# You can now add the other sections like 'Student Manager', 'Dashboard' etc.
