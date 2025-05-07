
# dance_school_app_combined.py

import streamlit as st
import gspread
import pandas as pd
from oauth2client.service_account import ServiceAccountCredentials
from datetime import datetime, date, timedelta
import calendar
import os
from docx import Document
from docx.shared import Inches

# Streamlit secrets setup for Google Sheets
scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
creds_dict = st.secrets["gcp_service_account"]
creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
client = gspread.authorize(creds)

# Main navigation
st.sidebar.title("Dance School App")
selection = st.sidebar.radio("Go to", [
    "Invoice Generator",
    "Student Manager",
    "Dashboard",
    "Attendance Register",
    "Accounts Package"
])

# 1. INVOICE GENERATOR
if selection == "Invoice Generator":
    exec(open("invoice_generator_block.py").read())

# 2. STUDENT MANAGER
elif selection == "Student Manager":
    exec(open("student_manager_block.py").read())

# 3. DASHBOARD
elif selection == "Dashboard":
    exec(open("dashboard_block.py").read())

# 4. ATTENDANCE REGISTER
elif selection == "Attendance Register":
    exec(open("attendance_register_block.py").read())

# 5. ACCOUNTS PACKAGE
elif selection == "Accounts Package":
    exec(open("accounts_package_block.py").read())

# Helper function for generating Word invoices
def generate_invoice_doc(student_name, date_from, date_to, class_list, extras, total):
    doc = Document()
    doc.add_heading("INVOICE", level=0)
    doc.add_paragraph(f"Student: {student_name}")
    doc.add_paragraph(f"Invoice Period: {date_from} to {date_to}")

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Description'
    hdr_cells[1].text = 'Amount'

    for item in class_list:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = ""

    for item in extras:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = ""

    row_cells = table.add_row().cells
    row_cells[0].text = 'Total'
    row_cells[1].text = f"£{total:.2f}"

    file_path = f"{student_name}_invoice.docx"
    doc.save(file_path)
    return file_path
